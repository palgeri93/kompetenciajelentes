import io
import re
import textwrap
from typing import Any

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Kompetenciamérés jelentésgenerátor", layout="wide")

ANGOL_SZINT_MAP = {"pre-a1": 1, "a1": 2, "a2": 3, "b1": 4, "b2": 5, "c1": 6}
ANGOL_SZINT_LABELS = {
    0: "0. szint - pre-A1 alatti szint",
    1: "1. szint - pre-A1 szint",
    2: "2. szint - A1 szint",
    3: "3. szint - A2 szint",
    4: "4. szint - B1 szint",
    5: "5. szint - B2 szint",
    6: "6. szint - C1 szint vagy fölötte",
}
BASE_COLUMNS = ["Évfolyam", "Tanulócsoportok", "Mérési azonosító", "Név"]
DEFAULT_AREAS = [
    "Szövegértés", "Matematika", "Természettudomány", "Angol nyelv",
    "Német nyelv", "Digitális kultúra", "Történelem"
]
CHANGE_ORDER = [
    "Nincs mindkét eredmény",
    "Jelentős -",
    "Mérsékelt -",
    "Elhanyagolható",
    "Mérsékelt +",
    "Jelentős +",
]
CHANGE_COLORS = {
    "Nincs mindkét eredmény": "#bdbdbd",
    "Jelentős -": "#8B0000",
    "Mérsékelt -": "#FF6347",
    "Elhanyagolható": "#9e9e9e",
    "Mérsékelt +": "#90EE90",
    "Jelentős +": "#006400",
}
NO_RESULT_LEVEL = -1
LEVEL_COLOR_MAP = {
    NO_RESULT_LEVEL: "#bdbdbd",  # nincs eredmény
    0: "#ffe08a",                # 0. szint - külön szín, nem szürke
    1: "#dff7d8",
    2: "#b9efad",
    3: "#62d66f",
    4: "#3fc76b",
    5: "#22bf73",
    6: "#17aa67",
    7: "#0c8f58",
}
LEVEL_COLORS = [LEVEL_COLOR_MAP[i] for i in range(0, 8)]


def clean_text(value: Any) -> str:
    if pd.isna(value):
        return ""
    return " ".join(str(value).strip().split())


def normalize_name_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    if "Név" not in df.columns and "NEVEK" in df.columns:
        df = df.rename(columns={"NEVEK": "Név"})
    return df


def decode_angol_szint(value: object) -> float:
    text = clean_text(value).lower().replace("–", "-")
    if not text or text in {"nan", "none", "nincs eredmény"}:
        return np.nan
    if "pre-a1" in text and "alatti" in text:
        return 0.0
    for key, level in ANGOL_SZINT_MAP.items():
        if key in text:
            return float(level)
    return np.nan


def decode_numeric_level(value: object) -> float:
    text = clean_text(value).lower()
    if not text or text in {"nan", "none", "nincs eredmény"}:
        return np.nan
    match = re.search(r"(\d+)", text)
    return float(match.group(1)) if match else np.nan


def categorize_change(value: float) -> tuple[str, str]:
    if value <= -100:
        return "#8B0000", "Jelentős romlás"
    if -99 <= value <= -40:
        return "#FF6347", "Mérsékelt romlás"
    if -39 <= value < 0:
        return "gray", "Elhanyagolható romlás"
    if 0 <= value <= 39:
        return "gray", "Elhanyagolható javulás"
    if 40 <= value <= 100:
        return "#90EE90", "Mérsékelt javulás"
    if value > 100:
        return "#006400", "Jelentős javulás"
    return "black", "Ismeretlen"


def short_change_category(row: pd.Series) -> str:
    p1 = pd.to_numeric(row.get("Képességpont_1"), errors="coerce")
    p2 = pd.to_numeric(row.get("Képességpont_2"), errors="coerce")
    if pd.isna(p1) and pd.isna(p2):
        return "Nincs mindkét eredmény"
    value = float(row.get("Képességpont változás", 0) or 0)
    if value <= -100:
        return "Jelentős -"
    if -99 <= value <= -40:
        return "Mérsékelt -"
    if -39 <= value <= 39:
        return "Elhanyagolható"
    if 40 <= value <= 100:
        return "Mérsékelt +"
    if value > 100:
        return "Jelentős +"
    return "Elhanyagolható"


def get_sheet_names(uploaded_file) -> list[str]:
    uploaded_file.seek(0)
    names = pd.ExcelFile(uploaded_file).sheet_names
    uploaded_file.seek(0)
    return names


def detect_areas(uploaded_file, sheet_name: str, area_row_index: int = 0) -> dict[str, int]:
    uploaded_file.seek(0)
    preview = pd.read_excel(uploaded_file, sheet_name=sheet_name, header=None, nrows=3)
    uploaded_file.seek(0)
    areas: dict[str, int] = {}
    for col_index, value in enumerate(preview.iloc[area_row_index].tolist()):
        name = clean_text(value)
        if name and name.lower() not in {"nan", "none"}:
            areas[name] = col_index
    return areas



def normalize_area_label(value: object) -> str:
    text = clean_text(value).lower()
    text = text.replace("angol nyelv", "angol").replace("német nyelv", "német")
    text = re.sub(r"[^a-záéíóöőúüű0-9]+", "", text)
    return text


def find_column(columns: list[object], required_words: list[str]) -> object | None:
    for col in columns:
        text = clean_text(col).lower()
        if all(word in text for word in required_words):
            return col
    return None


def read_level_thresholds(uploaded_file) -> pd.DataFrame:
    """Beolvassa a képességszint-ponthatárokat bármelyik Excel-fülről.

    Elvárt oszlopok rugalmasan felismerve:
    Kompetenciaterület / Szint / Alsó ponthatár.
    """
    uploaded_file.seek(0)
    xls = pd.ExcelFile(uploaded_file)
    uploaded_file.seek(0)
    candidates = []
    for sheet in xls.sheet_names:
        uploaded_file.seek(0)
        raw = pd.read_excel(uploaded_file, sheet_name=sheet)
        uploaded_file.seek(0)
        if raw.empty:
            continue
        area_col = find_column(list(raw.columns), ["kompetencia"]) or find_column(list(raw.columns), ["terület"])
        level_col = find_column(list(raw.columns), ["szint"])
        threshold_col = find_column(list(raw.columns), ["alsó", "ponthatár"]) or find_column(list(raw.columns), ["ponthatár"])
        if area_col is None or level_col is None or threshold_col is None:
            continue
        df_thr = raw[[area_col, level_col, threshold_col]].copy()
        df_thr.columns = ["Kompetenciaterület", "Szint", "Alsó ponthatár"]
        df_thr["Kompetenciaterület"] = df_thr["Kompetenciaterület"].map(clean_text)
        df_thr["Szint"] = pd.to_numeric(df_thr["Szint"], errors="coerce")
        df_thr["Alsó ponthatár"] = pd.to_numeric(df_thr["Alsó ponthatár"], errors="coerce")
        df_thr = df_thr.dropna(subset=["Kompetenciaterület", "Szint", "Alsó ponthatár"])
        if not df_thr.empty:
            candidates.append(df_thr)
    if not candidates:
        return pd.DataFrame(columns=["Kompetenciaterület", "Szint", "Alsó ponthatár"])
    return pd.concat(candidates, ignore_index=True).drop_duplicates()


def thresholds_for_area(thresholds_df: pd.DataFrame, terulet: str) -> pd.DataFrame:
    if thresholds_df is None or thresholds_df.empty:
        return pd.DataFrame(columns=["Képességszint", "Képességpont alsó ponthatár"])
    target = normalize_area_label(terulet)
    df_thr = thresholds_df.copy()
    df_thr["_norm"] = df_thr["Kompetenciaterület"].map(normalize_area_label)
    area_options = df_thr["_norm"].dropna().unique().tolist()
    matches = [x for x in area_options if x and (x in target or target in x)]
    if not matches:
        return pd.DataFrame(columns=["Képességszint", "Képességpont alsó ponthatár"])
    out = df_thr[df_thr["_norm"].isin(matches)].copy()
    out = out.sort_values("Szint", ascending=False)
    return pd.DataFrame({
        "Képességszint": out["Szint"].astype(int).map(lambda x: f"{x}. szint"),
        "Képességpont alsó ponthatár": out["Alsó ponthatár"].astype(int),
    })

def read_selected_area(uploaded_file, sheet_name: str, area_name: str, header_row_index: int = 2) -> pd.DataFrame:
    """
    Egy kiválasztott mérési terület beolvasása oszloppozíciók alapján.

    Fontos javítás: az Excelben több mérési terület alatt is ugyanazok az oszlopnevek
    szerepelnek (Képességpont, Képességszint stb.). Pandasban a név szerinti
    kiválasztás ilyenkor az összes azonos nevű oszlopot visszaadhatja, ezért itt
    kizárólag oszloppozícióval dolgozunk. Ez javítja a 4. ábra hibáját is, ahol
    korábban 100+ hamis „szint” jelent meg a jelmagyarázatban.
    """
    uploaded_file.seek(0)
    full_headerless = pd.read_excel(uploaded_file, sheet_name=sheet_name, header=None)
    uploaded_file.seek(0)
    area_map = detect_areas(uploaded_file, sheet_name, 0)
    uploaded_file.seek(0)
    if area_name not in area_map:
        raise ValueError(f"Nem találom ezt a mérési területet: {area_name}")

    start = int(area_map[area_name])
    next_starts = sorted(idx for name, idx in area_map.items() if idx > start)
    end = int(next_starts[0]) if next_starts else full_headerless.shape[1]

    headers = [clean_text(x) for x in full_headerless.iloc[header_row_index].tolist()]
    data = full_headerless.iloc[header_row_index + 1 :].reset_index(drop=True)

    # Alapadatok: az első négy oszlop fixen Évfolyam, Tanulócsoportok, Mérési azonosító, Név/NEVEK.
    result = pd.DataFrame()
    for pos, out_name in [(0, "Évfolyam"), (1, "Tanulócsoportok"), (2, "Mérési azonosító"), (3, "Név")]:
        if pos < data.shape[1]:
            result[out_name] = data.iloc[:, pos]

    # A kiválasztott mérési terület blokkjának oszlopai.
    area = data.iloc[:, start:end].copy()
    area_headers = headers[start:end]
    if area.shape[1] < 4:
        raise ValueError("A kiválasztott területnél nincs elég adat a diagramhoz.")

    # A jelentésekben az első négy területi oszlop: előző tanévi pont/szint, előzetes pont/szint.
    result["Képességpont_1"] = area.iloc[:, 0]
    result["Képességszint"] = area.iloc[:, 1]
    result["Képességpont_2"] = area.iloc[:, 2]
    result["Képességszint.1"] = area.iloc[:, 3]

    change_col = None
    for idx, col_name in enumerate(area_headers):
        col_text = str(col_name).lower()
        if "változás" in col_text and "%" not in col_text:
            change_col = area.iloc[:, idx]
            break
    if change_col is None:
        result["Képességpont változás"] = pd.to_numeric(result["Képességpont_2"], errors="coerce") - pd.to_numeric(result["Képességpont_1"], errors="coerce")
    else:
        result["Képességpont változás"] = change_col
    return result


def get_classes(df: pd.DataFrame) -> list[str]:
    if "Tanulócsoportok" not in df.columns:
        return ["Összes tanuló"]
    classes = sorted([x for x in df["Tanulócsoportok"].dropna().map(clean_text).unique() if x])
    return ["Összes tanuló"] + classes


def prepare_dataframe(df: pd.DataFrame, selected_class: str, angol: bool) -> pd.DataFrame:
    df = df.copy()
    if selected_class != "Összes tanuló" and "Tanulócsoportok" in df.columns:
        df = df[df["Tanulócsoportok"].map(clean_text).eq(selected_class)]
    if angol:
        df["Szint_1"] = df["Képességszint"].apply(decode_angol_szint)
        df["Szint_2"] = df["Képességszint.1"].apply(decode_angol_szint)
    else:
        df["Szint_1"] = df["Képességszint"].apply(decode_numeric_level)
        df["Szint_2"] = df["Képességszint.1"].apply(decode_numeric_level)
    df["Képességpont változás"] = pd.to_numeric(df["Képességpont változás"], errors="coerce").fillna(0)
    df["Név"] = df["Név"].astype(str).apply(lambda x: " ".join(x.split()))
    df = df[df["Név"].str.lower().ne("nan") & df["Név"].ne("")]
    df["Változás kategória"] = df.apply(short_change_category, axis=1)
    return df.sort_values(by="Képességpont változás")


def change_summary(df: pd.DataFrame, group_label: str) -> pd.DataFrame:
    counts = df["Változás kategória"].value_counts().reindex(CHANGE_ORDER, fill_value=0)
    total = max(len(df), 1)
    percent = (counts / total * 100).round(1)
    return pd.DataFrame([
        [group_label] + counts.astype(int).tolist(),
        ["Tanulók aránya"] + [f"{p:.1f}%".replace(".", ",") for p in percent.tolist()],
        ["Összesen"] + counts.astype(int).tolist(),
        ["Tanulók aránya"] + [f"{p:.1f}%".replace(".", ",") for p in percent.tolist()],
    ], columns=["Tanulócsoportok"] + CHANGE_ORDER)




def change_name_groups(df: pd.DataFrame) -> dict[str, list[str]]:
    """Név szerinti bontás a változás mértéke szerint, a Word-jelentéshez is másolható formában."""
    mapping = {
        "Nincs mindkét eredmény": "Nincs mindkét eredmény",
        "Jelentős -": "Jelentős romlás",
        "Mérsékelt -": "Mérsékelt romlás",
        "Elhanyagolható": "Elhanyagolható változás",
        "Mérsékelt +": "Mérsékelt javulás",
        "Jelentős +": "Jelentős javulás",
    }
    groups = {label: [] for label in mapping.values()}
    for _, row in df.sort_values("Képességpont változás").iterrows():
        short = row.get("Változás kategória", "Elhanyagolható")
        label = mapping.get(short, "Elhanyagolható változás")
        name = clean_text(row.get("Név", ""))
        if name:
            groups[label].append(name)
    return groups


def change_names_text(df: pd.DataFrame) -> str:
    groups = change_name_groups(df)
    lines = []
    templates = [
        ("Jelentős romlás", "Jelentős romlás {count} tanulónál történt: {names}"),
        ("Mérsékelt romlás", "Mérsékelt romlás {count} főnél történt: {names}"),
        ("Elhanyagolható változás", "Elhanyagolható a változás {count} tanuló esetében: {names}"),
        ("Mérsékelt javulás", "Mérsékelt javulást mutat {count} tanuló: {names}"),
        ("Jelentős javulás", "Jelentős mértékű javulást mutat {count} fő: {names}"),
        ("Nincs mindkét eredmény", "Nincs mindkét eredménye {count} tanulónak: {names}"),
    ]
    for label, template in templates:
        names = groups.get(label, [])
        if not names and label == "Nincs mindkét eredmény":
            continue
        lines.append(template.format(count=len(names), names=", ".join(names) if names else "-"))
    return "\n".join(lines)


def level_distribution(df: pd.DataFrame, column: str, max_level: int) -> dict[int, int]:
    vals = pd.to_numeric(df[column], errors="coerce")
    dist = {NO_RESULT_LEVEL: int(vals.isna().sum())}
    filled = vals.dropna().astype(int)
    dist.update({level: int((filled == level).sum()) for level in range(0, max_level + 1)})
    return dist


def level_label(level: int) -> str:
    if level == NO_RESULT_LEVEL:
        return "Nincs eredmény"
    return f"{level}. szint"


def level_color(level: int) -> str:
    return LEVEL_COLOR_MAP.get(int(level), "#cccccc")


def add_change_chart(pdf: PdfPages, df: pd.DataFrame, osztaly: str, terulet: str) -> None:
    if df.empty or df["Képességpont változás"].abs().sum() <= 0:
        return
    colors, _ = zip(*[categorize_change(v) for v in df["Képességpont változás"]])
    plt.figure(figsize=(14, 6))
    bars = plt.bar(df["Név"], df["Képességpont változás"], color=colors)
    plt.axhline(0, color="black", linewidth=0.8)
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width() / 2, height, f"{int(height)}", ha="center", va="bottom" if height >= 0 else "top", fontsize=8)
    plt.title(f"{osztaly} – {terulet} – Képességpont változás")
    plt.xlabel("Tanuló neve")
    plt.ylabel("Képességpont változás")
    plt.xticks(rotation=45, ha="right")
    legend_labels = {"Jelentős romlás": "#8B0000", "Mérsékelt romlás": "#FF6347", "Elhanyagolható romlás/javulás": "gray", "Mérsékelt javulás": "#90EE90", "Jelentős javulás": "#006400"}
    patches = [plt.Line2D([0], [0], color=color, lw=8, label=label) for label, color in legend_labels.items()]
    plt.legend(handles=patches, title="Kategóriák")
    plt.subplots_adjust(bottom=0.25)
    plt.tight_layout()
    pdf.savefig()
    plt.close()


def add_change_summary_table(pdf: PdfPages, df: pd.DataFrame, osztaly: str, terulet: str) -> None:
    table_df = change_summary(df, osztaly)
    fig, ax = plt.subplots(figsize=(14, 7.2))
    ax.axis("off")
    ax.set_title("Változás mértéke az elemzésbe bevont csoportok esetében", loc="left", fontsize=14, fontweight="bold", pad=18)
    ax.text(0.5, 0.90, terulet, ha="center", va="center", fontsize=11, fontweight="bold", transform=ax.transAxes)
    table = ax.table(cellText=table_df.values, colLabels=table_df.columns, cellLoc="center", bbox=[0.02, 0.48, 0.96, 0.34])
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.5)
    for (row, col), cell in table.get_celld().items():
        cell.set_linewidth(0.3)
        if row == 0:
            cell.set_text_props(fontweight="bold")
            cell.set_facecolor("#f3f3f3")
        if row in [1, 3]:
            cell.set_facecolor("#eeeeee")

    # Név szerinti felsorolás ugyanarra az oldalra kerül, a Word-jelentés 3. ábra alatti részéhez hasonlóan.
    y = 0.39
    for raw_line in change_names_text(df).splitlines():
        wrapped = textwrap.wrap(raw_line, width=125) or [raw_line]
        for line in wrapped:
            ax.text(0.02, y, line, ha="left", va="top", fontsize=10.5, transform=ax.transAxes)
            y -= 0.055
        y -= 0.012
    plt.tight_layout()
    pdf.savefig()
    plt.close()

def add_level_distribution_chart(pdf: PdfPages, df: pd.DataFrame, osztaly: str, terulet: str, angol: bool) -> None:
    if df.empty:
        return
    max_level = int(6 if angol else 7)
    rows = [("2024/2025-ös tanév", "Szint_1"), ("2025/2026-os tanév előzetes eredmény", "Szint_2")]
    fig, ax = plt.subplots(figsize=(14, 6.0))
    y_positions = np.arange(len(rows))
    total = max(len(df), 1)
    present_levels = set()

    for y, (label, col) in zip(y_positions, rows):
        left = 0.0
        dist = level_distribution(df, col, max_level)
        for level in [NO_RESULT_LEVEL] + list(range(0, max_level + 1)):
            count = dist[level]
            if count == 0:
                continue
            present_levels.add(level)
            pct = count / total * 100
            color = level_color(level)
            ax.barh(y, pct, left=left, height=0.46, color=color, edgecolor="white")
            label_text = f"{pct:.1f} %".replace(".", ",")

            # Minden nem nulla szegmens kap százalékfeliratot.
            # A nagyon keskeny szegmenseknél a feliratot a sávon kívülre tesszük,
            # különben a szöveg vagy lemaradna, vagy olvashatatlanul összecsúszna.
            center = left + pct / 2
            if pct >= 7:
                ax.text(center, y, label_text, ha="center", va="center", fontsize=10, fontweight="bold")
            elif pct >= 3:
                ax.text(center, y, label_text, ha="center", va="center", fontsize=8.5, fontweight="bold")
            else:
                offset_y = -0.34 if y == 0 else 0.34
                label_x = min(max(center, 2.0), 98.0)
                ax.annotate(
                    label_text,
                    xy=(center, y),
                    xytext=(label_x, y + offset_y),
                    textcoords="data",
                    ha="center",
                    va="center",
                    fontsize=8,
                    fontweight="bold",
                    arrowprops={"arrowstyle": "-", "linewidth": 0.6, "color": "black", "shrinkA": 0, "shrinkB": 0},
                    clip_on=False,
                )
            left += pct

    ax.set_yticks(y_positions)
    ax.set_yticklabels([r[0] for r in rows])
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("Tanulók aránya")
    ax.set_title(f"{osztaly} – {terulet} – Mérési szintek megoszlása", fontsize=14, fontweight="bold")
    ax.set_xticks([0, 20, 40, 60, 80, 100])
    ax.set_xticklabels([f"{x},0 %" for x in [0, 20, 40, 60, 80, 100]])

    legend_levels = [level for level in [NO_RESULT_LEVEL] + list(range(0, max_level + 1)) if level in present_levels]
    handles = [
        plt.Line2D([0], [0], color=level_color(level), lw=8,
                   label=level_label(level))
        for level in legend_levels
    ]
    if handles:
        ax.legend(handles=handles, loc="lower center", bbox_to_anchor=(0.5, -0.30), ncol=min(len(handles), 8))
    ax.grid(axis="x", alpha=0.25)
    fig.subplots_adjust(left=0.23, bottom=0.25, right=0.98, top=0.86)
    pdf.savefig(fig)
    plt.close(fig)


def add_level_chart(pdf: PdfPages, df: pd.DataFrame, alapszint: int, osztaly: str, terulet: str, angol: bool) -> None:
    if df.empty:
        return
    x = np.arange(len(df))
    width = 0.35
    fig, ax = plt.subplots(figsize=(14, 6))
    ax.bar(x - width / 2, df["Szint_1"], width, label="2024/2025")
    ax.bar(x + width / 2, df["Szint_2"], width, label="2025/2026 előzetes")

    max_data_level = int(max(df["Szint_1"].max(), df["Szint_2"].max(), alapszint, 6 if angol else 7))
    ax.set_ylim(0, max_data_level + 0.8)
    if alapszint > 0:
        ax.axhline(alapszint, color="red", linestyle="--", linewidth=2, label=f"Alapszint: {alapszint}", zorder=3)

    ax.set_title(f"{osztaly} – {terulet} – Képességszint változás")
    ax.set_xlabel("Tanuló neve")
    ax.set_ylabel("Képességszint")
    ax.set_xticks(x)
    ax.set_xticklabels(df["Név"], rotation=45, ha="right")
    if angol:
        ticks = sorted(ANGOL_SZINT_LABELS.keys())
        ax.set_yticks(ticks)
        ax.set_yticklabels([ANGOL_SZINT_LABELS[y] for y in ticks])
    else:
        ax.set_yticks(np.arange(0, max_data_level + 1, 1))
    ax.legend(loc="upper right")
    ax.grid(axis="y", alpha=0.2)
    fig.subplots_adjust(left=0.08, bottom=0.28, right=0.98, top=0.90)
    pdf.savefig(fig)
    plt.close(fig)



def make_change_chart_fig(df: pd.DataFrame, osztaly: str, terulet: str):
    if df.empty or df["Képességpont változás"].abs().sum() <= 0:
        return None
    colors, _ = zip(*[categorize_change(v) for v in df["Képességpont változás"]])
    fig, ax = plt.subplots(figsize=(14, 6))
    bars = ax.bar(df["Név"], df["Képességpont változás"], color=colors)
    ax.axhline(0, color="black", linewidth=0.8)
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, f"{int(height)}", ha="center", va="bottom" if height >= 0 else "top", fontsize=8)
    ax.set_title(f"{osztaly} – {terulet} – Képességpont változás")
    ax.set_xlabel("Tanuló neve")
    ax.set_ylabel("Képességpont változás")
    ax.set_xticklabels(df["Név"], rotation=45, ha="right")
    legend_labels = {"Jelentős romlás": "#8B0000", "Mérsékelt romlás": "#FF6347", "Elhanyagolható romlás/javulás": "gray", "Mérsékelt javulás": "#90EE90", "Jelentős javulás": "#006400"}
    patches = [plt.Line2D([0], [0], color=color, lw=8, label=label) for label, color in legend_labels.items()]
    ax.legend(handles=patches, title="Kategóriák")
    fig.subplots_adjust(bottom=0.25)
    fig.tight_layout()
    return fig


def make_change_summary_fig(df: pd.DataFrame, osztaly: str, terulet: str):
    table_df = change_summary(df, osztaly)
    fig, ax = plt.subplots(figsize=(14, 7.2))
    ax.axis("off")
    ax.set_title("Változás mértéke az elemzésbe bevont csoportok esetében", loc="left", fontsize=14, fontweight="bold", pad=18)
    ax.text(0.5, 0.90, terulet, ha="center", va="center", fontsize=11, fontweight="bold", transform=ax.transAxes)
    table = ax.table(cellText=table_df.values, colLabels=table_df.columns, cellLoc="center", bbox=[0.02, 0.48, 0.96, 0.34])
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.5)
    for (row, col), cell in table.get_celld().items():
        cell.set_linewidth(0.3)
        if row == 0:
            cell.set_text_props(fontweight="bold")
            cell.set_facecolor("#f3f3f3")
        if row in [1, 3]:
            cell.set_facecolor("#eeeeee")
    y = 0.39
    for raw_line in change_names_text(df).splitlines():
        wrapped = textwrap.wrap(raw_line, width=125) or [raw_line]
        for line in wrapped:
            ax.text(0.02, y, line, ha="left", va="top", fontsize=10.5, transform=ax.transAxes)
            y -= 0.055
        y -= 0.012
    fig.tight_layout()
    return fig


def make_level_distribution_fig(df: pd.DataFrame, osztaly: str, terulet: str, angol: bool):
    if df.empty:
        return None
    max_level = int(6 if angol else 7)
    rows = [("2024/2025-ös tanév", "Szint_1"), ("2025/2026-os tanév előzetes eredmény", "Szint_2")]
    fig, ax = plt.subplots(figsize=(14, 6.0))
    y_positions = np.arange(len(rows))
    total = max(len(df), 1)
    present_levels = set()

    for y, (label, col) in zip(y_positions, rows):
        left = 0.0
        dist = level_distribution(df, col, max_level)
        for level in [NO_RESULT_LEVEL] + list(range(0, max_level + 1)):
            count = dist[level]
            if count == 0:
                continue
            present_levels.add(level)
            pct = count / total * 100
            color = level_color(level)
            ax.barh(y, pct, left=left, height=0.46, color=color, edgecolor="white")
            label_text = f"{pct:.1f} %".replace(".", ",")
            center = left + pct / 2
            if pct >= 7:
                ax.text(center, y, label_text, ha="center", va="center", fontsize=10, fontweight="bold")
            elif pct >= 3:
                ax.text(center, y, label_text, ha="center", va="center", fontsize=8.5, fontweight="bold")
            else:
                offset_y = -0.34 if y == 0 else 0.34
                label_x = min(max(center, 2.0), 98.0)
                ax.annotate(
                    label_text,
                    xy=(center, y),
                    xytext=(label_x, y + offset_y),
                    textcoords="data",
                    ha="center",
                    va="center",
                    fontsize=8,
                    fontweight="bold",
                    arrowprops={"arrowstyle": "-", "linewidth": 0.6, "color": "black", "shrinkA": 0, "shrinkB": 0},
                    clip_on=False,
                )
            left += pct

    ax.set_yticks(y_positions)
    ax.set_yticklabels([r[0] for r in rows])
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("Tanulók aránya")
    ax.set_title(f"{osztaly} – {terulet} – Mérési szintek megoszlása", fontsize=14, fontweight="bold")
    ax.set_xticks([0, 20, 40, 60, 80, 100])
    ax.set_xticklabels([f"{x},0 %" for x in [0, 20, 40, 60, 80, 100]])
    legend_levels = [level for level in [NO_RESULT_LEVEL] + list(range(0, max_level + 1)) if level in present_levels]
    handles = [
        plt.Line2D([0], [0], color=level_color(level), lw=8,
                   label=level_label(level))
        for level in legend_levels
    ]
    if handles:
        ax.legend(handles=handles, loc="lower center", bbox_to_anchor=(0.5, -0.30), ncol=min(len(handles), 8))
    ax.grid(axis="x", alpha=0.25)
    fig.subplots_adjust(left=0.23, bottom=0.25, right=0.98, top=0.86)
    return fig


def make_level_chart_fig(df: pd.DataFrame, alapszint: int, osztaly: str, terulet: str, angol: bool):
    if df.empty:
        return None
    x = np.arange(len(df))
    width = 0.35
    fig, ax = plt.subplots(figsize=(14, 6))
    ax.bar(x - width / 2, df["Szint_1"], width, label="2024/2025")
    ax.bar(x + width / 2, df["Szint_2"], width, label="2025/2026 előzetes")
    max_data_level = int(max(df["Szint_1"].max(), df["Szint_2"].max(), alapszint, 6 if angol else 7))
    ax.set_ylim(0, max_data_level + 0.8)
    if alapszint > 0:
        ax.axhline(alapszint, color="red", linestyle="--", linewidth=2, label=f"Alapszint: {alapszint}", zorder=3)
    ax.set_title(f"{osztaly} – {terulet} – Képességszint változás")
    ax.set_xlabel("Tanuló neve")
    ax.set_ylabel("Képességszint")
    ax.set_xticks(x)
    ax.set_xticklabels(df["Név"], rotation=45, ha="right")
    if angol:
        ticks = sorted(ANGOL_SZINT_LABELS.keys())
        ax.set_yticks(ticks)
        ax.set_yticklabels([ANGOL_SZINT_LABELS[y] for y in ticks])
    else:
        ax.set_yticks(np.arange(0, max_data_level + 1, 1))
    ax.legend(loc="upper right")
    ax.grid(axis="y", alpha=0.2)
    fig.subplots_adjust(left=0.08, bottom=0.28, right=0.98, top=0.90)
    return fig



def fmt_num(value: float, digits: int = 0) -> str:
    if pd.isna(value):
        return "-"
    if digits == 0:
        return f"{float(value):.0f}".replace(".", ",")
    return f"{float(value):.{digits}f}".replace(".", ",")


def fmt_pct(value: float, digits: int = 1) -> str:
    if pd.isna(value):
        return "-"
    return f"{float(value):.{digits}f}%".replace(".", ",")


def direction_word(value: float, noun: str = "változás") -> str:
    if value < 0:
        return "romlás" if noun == "változás" else "romlást"
    if value > 0:
        return "javulás" if noun == "változás" else "javulást"
    return "változás" if noun == "változás" else "változást"


def count_word(n: int) -> str:
    words = {0: "Nulla", 1: "Egy", 2: "Két", 3: "Három", 4: "Négy", 5: "Öt", 6: "Hat", 7: "Hét", 8: "Nyolc", 9: "Kilenc", 10: "Tíz"}
    return words.get(int(n), str(int(n)))


def add_df_table_to_doc(doc: Document, df_table: pd.DataFrame, max_rows: int | None = None) -> None:
    out = df_table.copy()
    if max_rows is not None:
        out = out.head(max_rows)
    table = doc.add_table(rows=1, cols=len(out.columns))
    table.style = "Table Grid"
    for i, col in enumerate(out.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in out.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            cells[i].text = clean_text(value)


def fig_to_png_bytes(fig) -> io.BytesIO:
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    buffer.seek(0)
    return buffer


def report_metrics(df: pd.DataFrame) -> dict[str, Any]:
    p1 = pd.to_numeric(df["Képességpont_1"], errors="coerce")
    p2 = pd.to_numeric(df["Képességpont_2"], errors="coerce")
    s1 = pd.to_numeric(df["Szint_1"], errors="coerce")
    s2 = pd.to_numeric(df["Szint_2"], errors="coerce")
    both_points = p1.notna() & p2.notna()
    avg_p1 = float(p1[both_points].mean()) if both_points.any() else 0.0
    avg_p2 = float(p2[both_points].mean()) if both_points.any() else 0.0
    point_diff = avg_p2 - avg_p1
    point_pct = (point_diff / avg_p1 * 100) if avg_p1 else 0.0
    both_levels = s1.notna() & s2.notna()
    avg_s1 = float(s1[both_levels].mean()) if both_levels.any() else 0.0
    avg_s2 = float(s2[both_levels].mean()) if both_levels.any() else 0.0
    level_diff = avg_s2 - avg_s1
    level_pct = (level_diff / avg_s1 * 100) if avg_s1 else 0.0
    counts = df["Változás kategória"].value_counts().reindex(CHANGE_ORDER, fill_value=0).astype(int)
    return {
        "avg_p1": avg_p1,
        "avg_p2": avg_p2,
        "point_diff": point_diff,
        "point_pct": point_pct,
        "avg_s1": avg_s1,
        "avg_s2": avg_s2,
        "level_diff": level_diff,
        "level_pct": level_pct,
        "counts": counts,
        "negative": int((pd.to_numeric(df["Képességpont változás"], errors="coerce") < 0).sum()),
        "positive": int((pd.to_numeric(df["Képességpont változás"], errors="coerce") > 0).sum()),
        "moderate_or_significant_positive": int(counts.get("Mérsékelt +", 0) + counts.get("Jelentős +", 0)),
    }


def level_sentence(df: pd.DataFrame, col: str, max_level: int, year_label: str) -> str:
    dist = level_distribution(df, col, max_level)
    total = max(len(df), 1)
    parts = []
    for level in [NO_RESULT_LEVEL] + list(range(0, max_level + 1)):
        count = dist[level]
        if count:
            label = "nincs eredménye" if level == NO_RESULT_LEVEL else f"{level}. szinten"
            parts.append(f"{label} {fmt_pct(count / total * 100, 1)}-a ({count} fő)")
    return f"A {year_label} tanévben a csoport tanulóinak " + ", ".join(parts) + " teljesített."


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def generate_word_report(
    df: pd.DataFrame,
    alapszint: int,
    osztaly: str,
    terulet: str,
    angol: bool,
    report_year: str,
    previous_year: str,
    current_year: str,
    institution_name: str,
    include_summary: bool = True,
    include_distribution: bool = True,
    thresholds_table: pd.DataFrame | None = None,
) -> bytes:
    """Szerkeszthető Word-jelentés automatikus számított szövegekkel, táblázatokkal és ábrákkal."""
    max_level = 6 if angol else 7
    metrics = report_metrics(df)
    counts = metrics["counts"]
    groups = change_name_groups(df)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{report_year} DIGITÁLIS ORSZÁGOS KOMPETENCIAMÉRÉS ELEMZÉSE ELŐZETES EREDMÉNYEK ALAPJÁN")
    run.bold = True
    run.font.size = None
    doc.add_paragraph(osztaly.upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(terulet.upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"A {institution_name} a {report_year} évi országos kompetenciamérésének értékelése az előzetes eredmények alapján. "
        f"{terulet} tantárgyból a {osztaly} csoportban {len(df)} tanuló szerepel az elemzésben. "
        f"A {previous_year}-es tanévben elért képességpontokat összehasonlítva a {current_year} tanév előzetes eredményeivel "
        f"{metrics['negative']} esetben mutatható ki képességpont csökkenés és {metrics['positive']} esetben javulás mérhető."
    )

    table_cols = ["Tanulócsoportok", "Mérési azonosító", "Név", "Képességpont_1", "Képességszint", "Képességpont_2", "Képességszint.1", "Képességpont változás"]
    data_table = df[[c for c in table_cols if c in df.columns]].copy()
    data_table = data_table.rename(columns={
        "Képességpont_1": f"Képességpont {previous_year}",
        "Képességszint": f"Képességszint {previous_year}",
        "Képességpont_2": f"Képességpont {current_year}",
        "Képességszint.1": f"Képességszint {current_year}",
    })
    data_table["Képességpont változás %-os formában"] = df.apply(
        lambda r: fmt_pct((pd.to_numeric(r["Képességpont változás"], errors="coerce") / pd.to_numeric(r["Képességpont_1"], errors="coerce") * 100) if pd.to_numeric(r["Képességpont_1"], errors="coerce") else 0, 2), axis=1
    )
    data_table["Változás mértéke"] = [categorize_change(v)[1].lower() for v in df["Képességpont változás"]]
    add_df_table_to_doc(doc, data_table)
    # Átlagsor évszámokkal, hogy a Word-jelentésben egyértelmű legyen,
    # melyik átlag melyik mérési tanévhez tartozik.
    doc.add_paragraph(
        f"ÁTLAG ({previous_year}): {fmt_num(metrics['avg_p1'], 2)} pont, {fmt_num(metrics['avg_s1'], 2)} szint; "
        f"ÁTLAG ({current_year} előzetes): {fmt_num(metrics['avg_p2'], 2)} pont, {fmt_num(metrics['avg_s2'], 2)} szint; "
        f"változás: {fmt_num(metrics['point_diff'], 2)} pont ({fmt_pct(metrics['point_pct'], 2)})."
    )

    add_heading(doc, "Képességpontok, képességpont változások", 1)
    doc.add_paragraph(
        f"Az osztály {previous_year} évi mérésen átlagosan {fmt_num(metrics['avg_p1'])} pontot ért el, a {current_year} évi mérés előzetes eredményei alapján "
        f"{fmt_num(metrics['avg_p2'])} pontot, így {fmt_num(metrics['point_diff'])} képességpontos {direction_word(metrics['point_diff'], 'tárgy')} mutat a csoport. "
        f"Százalékos formában {fmt_pct(abs(metrics['point_pct']), 2)} {direction_word(metrics['point_diff'])} figyelhető meg."
    )
    doc.add_paragraph(
        f"Az osztály átlag képességszintje {previous_year} tanévi mérés során {fmt_num(metrics['avg_s1'], 2)} volt, míg a {current_year} mérésen {fmt_num(metrics['avg_s2'], 2)}. "
        f"Százalékos formában ez {fmt_pct(abs(metrics['level_pct']), 1)} {direction_word(metrics['level_diff'])}nak tekinthető."
    )
    doc.add_paragraph(
        f"{count_word(int(counts.get('Jelentős -', 0)))} tanuló esetében mutatható ki jelentős romlás, {int(counts.get('Mérsékelt -', 0))} diák esetében mérsékelt romlás, "
        f"{metrics['positive']} tanuló esetében pozitív változás mutatható ki, melyből {int(counts.get('Mérsékelt +', 0))} tanuló mérsékelt javulást és "
        f"{int(counts.get('Jelentős +', 0))} tanuló jelentős javulást ért el. {count_word(int(counts.get('Elhanyagolható', 0)))} tanuló esetében elhanyagolható a változás. "
        f"A tanulók {fmt_pct(metrics['moderate_or_significant_positive'] / max(len(df), 1) * 100, 2)}-a mutatott legalább mérsékelt javulást."
    )

    fig = make_change_chart_fig(df, osztaly, terulet)
    if fig is not None:
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.9))
        doc.add_paragraph("1. ábra").alignment = WD_ALIGN_PARAGRAPH.CENTER
        plt.close(fig)

    if include_summary:
        add_heading(doc, "Változás mértéke az elemzésbe bevont csoportok esetében", 1)
        add_df_table_to_doc(doc, change_summary(df, osztaly))
        doc.add_paragraph("2. ábra").alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in change_names_text(df).splitlines():
            doc.add_paragraph(line)

    add_heading(doc, "Képességszintek, képességszint változások", 1)
    doc.add_paragraph(f"A {osztaly} csoportban a központilag megjelölt alap képességszint a {alapszint}. szint.")
    if include_distribution:
        fig = make_level_distribution_fig(df, osztaly, terulet, angol)
        if fig is not None:
            doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.9))
            doc.add_paragraph("3. ábra").alignment = WD_ALIGN_PARAGRAPH.CENTER
            plt.close(fig)
        doc.add_paragraph(level_sentence(df, "Szint_1", max_level, previous_year))
        doc.add_paragraph(level_sentence(df, "Szint_2", max_level, current_year))

    fig = make_level_chart_fig(df, alapszint, osztaly, terulet, angol)
    if fig is not None:
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.9))
        doc.add_paragraph("4. ábra").alignment = WD_ALIGN_PARAGRAPH.CENTER
        plt.close(fig)

    if thresholds_table is not None and not thresholds_table.empty:
        doc.add_paragraph(
            f"Az 5. ábra mutatja az adott képességszintek alsó ponthatárát a(z) {terulet} mérési területen."
        )
        add_df_table_to_doc(doc, thresholds_table)
        doc.add_paragraph("5. ábra").alignment = WD_ALIGN_PARAGRAPH.CENTER

    level_change = pd.to_numeric(df["Szint_2"], errors="coerce") - pd.to_numeric(df["Szint_1"], errors="coerce")
    improved_names = df.loc[level_change > 0, "Név"].tolist()
    worsened_names = df.loc[level_change < 0, "Név"].tolist()
    unchanged_names = df.loc[level_change == 0, "Név"].tolist()
    doc.add_paragraph(
        f"Mérési azonosítók alapján {len(improved_names)} tanuló ({', '.join(improved_names) if improved_names else '-'}) esetében volt mérhető képességszint javulás. "
        f"Képességszint romlás {len(worsened_names)} tanuló esetében volt ({', '.join(worsened_names) if worsened_names else '-'}). "
        f"{count_word(len(unchanged_names))} tanulónál nem jelentkezett változás a képességszintben."
    )

    current_levels = pd.to_numeric(df["Szint_2"], errors="coerce")
    missing_names = df.loc[current_levels.isna(), "Név"].tolist()
    if missing_names:
        doc.add_paragraph(f"Nincs eredmény: {', '.join(missing_names)}")
    for level in range(0, max_level + 1):
        names = df.loc[current_levels.dropna().reindex(df.index).fillna(-999).astype(int).eq(level), "Név"].tolist()
        if names:
            label = f"{level}. szinten teljesített tanulók"
            doc.add_paragraph(f"{label}: {', '.join(names)}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_thresholds_table_to_pdf(pdf: PdfPages, thresholds_table: pd.DataFrame, terulet: str) -> None:
    if thresholds_table is None or thresholds_table.empty:
        return
    fig, ax = plt.subplots(figsize=(8.5, 5.2))
    ax.axis("off")
    ax.set_title(
        f"{terulet} – képességszintek alsó ponthatára",
        loc="left",
        fontsize=14,
        fontweight="bold",
        pad=16,
    )
    table = ax.table(
        cellText=thresholds_table.values,
        colLabels=thresholds_table.columns,
        cellLoc="center",
        bbox=[0.05, 0.10, 0.90, 0.78],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(11)
    table.scale(1, 1.4)
    for (row, col), cell in table.get_celld().items():
        cell.set_linewidth(0.4)
        if row == 0:
            cell.set_text_props(fontweight="bold")
            cell.set_facecolor("#f3f3f3")
    fig.tight_layout()
    pdf.savefig(fig)
    plt.close(fig)

def generate_pdf(df: pd.DataFrame, alapszint: int, osztaly: str, terulet: str, angol: bool, include_summary: bool = True, include_distribution: bool = True, thresholds_table: pd.DataFrame | None = None) -> bytes:
    buffer = io.BytesIO()
    with PdfPages(buffer) as pdf:
        add_change_chart(pdf, df, osztaly, terulet)
        if include_summary:
            add_change_summary_table(pdf, df, osztaly, terulet)
        if include_distribution:
            add_level_distribution_chart(pdf, df, osztaly, terulet, angol)
        add_level_chart(pdf, df, alapszint, osztaly, terulet, angol)
        if thresholds_table is not None and not thresholds_table.empty:
            add_thresholds_table_to_pdf(pdf, thresholds_table, terulet)
    buffer.seek(0)
    return buffer.getvalue()


st.title("Kompetenciamérés jelentésgenerátor")
st.write("Tölts fel egy Excel fájlt, válassz mérési területet és osztályt, majd töltsd le a szerkeszthető Word-jelentést vagy a PDF diagramcsomagot.")

uploaded_file = st.file_uploader("Excel fájl feltöltése", type=["xlsx"])
if uploaded_file is None:
    st.info("Kezdéshez tölts fel egy .xlsx fájlt.")
    st.stop()

sheet_names = get_sheet_names(uploaded_file)
thresholds_df = read_level_thresholds(uploaded_file)
with st.sidebar:
    st.header("Beállítások")
    sheet_name = st.selectbox("Munkalap", sheet_names, index=sheet_names.index("Munka1") if "Munka1" in sheet_names else 0)
    header_row = st.number_input("Fejléc sor száma az Excelben", min_value=1, max_value=20, value=3)

try:
    area_map = detect_areas(uploaded_file, sheet_name)
    area_options = list(area_map.keys()) or DEFAULT_AREAS
except Exception as exc:
    st.error(f"Nem sikerült a mérési területek felismerése: {exc}")
    st.stop()

with st.sidebar:
    terulet = st.selectbox("Kompetenciamérési terület", area_options)

try:
    raw_df = read_selected_area(uploaded_file, sheet_name, terulet, int(header_row) - 1)
except Exception as exc:
    st.error(str(exc))
    st.stop()

class_options = get_classes(raw_df)
with st.sidebar:
    selected_class = st.selectbox("Osztály / tanulócsoport", class_options)
    default_base = 2 if "angol" not in terulet.lower() and "német" not in terulet.lower() else 3
    alapszint = st.number_input("Alapszint (0 = ne jelenjen meg)", min_value=0, max_value=10, value=default_base)
    auto_language = "angol" in terulet.lower() or "német" in terulet.lower()
    angol = st.checkbox("Nyelvi szintek használata (pre-A1, A1, A2...)", value=auto_language)
    include_summary = st.checkbox("3. ábra: változás mértéke táblázat", value=True)
    include_distribution = st.checkbox("4. ábra: mérési szintek halmozott sávdiagram", value=True)
    include_thresholds = st.checkbox("5. ábra: képességszintek alsó ponthatára", value=not thresholds_df.empty)
    safe_class = selected_class.replace(" ", "_").replace("/", "-")
    pdf_nev = st.text_input("PDF fájlnév", value=f"{terulet}_{safe_class}.pdf")
    docx_nev = st.text_input("Word fájlnév", value=f"{terulet}_{safe_class}_jelentes.docx")
    report_year = st.text_input("Jelentés éve", value="2026.")
    previous_year = st.text_input("Előző tanév", value="2024/2025")
    current_year = st.text_input("Aktuális/előzetes tanév", value="2025/2026")
    institution_name = st.text_input("Intézmény neve", value="Sulyok István Református Általános Iskola és AMI")

try:
    df = prepare_dataframe(raw_df, selected_class, angol)
except Exception as exc:
    st.error(str(exc))
    st.stop()

if df.empty:
    st.warning("A kiválasztott beállításokkal nincs megjeleníthető tanuló.")
    st.stop()

thresholds_table = thresholds_for_area(thresholds_df, terulet) if include_thresholds else pd.DataFrame()

st.subheader("Beolvasott adatok előnézete")
preview_cols = [c for c in ["Évfolyam", "Tanulócsoportok", "Mérési azonosító", "Név", "Képességpont_1", "Képességszint", "Képességpont_2", "Képességszint.1", "Képességpont változás", "Változás kategória", "Szint_1", "Szint_2"] if c in df.columns]
st.dataframe(df[preview_cols], use_container_width=True)
st.caption(f"{len(df)} tanuló • terület: {terulet} • osztály/tanulócsoport: {selected_class}")

col1, col2 = st.columns(2)
with col1:
    st.subheader("Változás mértéke összesítő")
    st.dataframe(change_summary(df, selected_class), use_container_width=True, hide_index=True)
with col2:
    st.subheader("Szintmegoszlás")
    max_level_preview = int(max(6 if angol else 7, df["Szint_1"].max(), df["Szint_2"].max()))
    level_rows = []
    for label, col in [("2024/2025", "Szint_1"), ("2025/2026 előzetes", "Szint_2")]:
        dist = level_distribution(df, col, max_level_preview)
        total = max(len(df), 1)
        row = {"Tanév": label}
        for level, count in dist.items():
            if count:
                col_label = "Nincs eredmény" if level == NO_RESULT_LEVEL else f"{level}. szint"
                row[col_label] = f"{count} fő ({count/total*100:.1f}%)".replace(".", ",")
        level_rows.append(row)
    st.dataframe(pd.DataFrame(level_rows).fillna("-"), use_container_width=True, hide_index=True)

if include_thresholds:
    st.subheader("Képességszintek alsó ponthatára")
    if thresholds_table.empty:
        st.warning("Nem találtam a kiválasztott mérési területhez ponthatár-táblát. Hozz létre egy új Excel-fület például ezekkel az oszlopokkal: Kompetenciaterület, Szint, Alsó ponthatár.")
    else:
        st.dataframe(thresholds_table, use_container_width=True, hide_index=True)

st.subheader("Név szerinti bontás a változás mértéke szerint")
st.write("Ez a rész a PDF-be is bekerül, és innen a Word-jelentésbe is bemásolható.")
st.text_area("Másolható szöveg", value=change_names_text(df), height=190)

st.subheader("Diagram előnézetek a PDF jelentésből")
st.write("Az alábbi ábrák ugyanazok, amelyek a PDF jelentésbe is bekerülnek.")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["1. ábra", "2. ábra", "3. ábra", "4. ábra", "5. ábra"] )
with tab1:
    fig = make_change_chart_fig(df, selected_class, terulet)
    if fig is None:
        st.info("Ehhez az ábrához nincs megjeleníthető képességpont-változás.")
    else:
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)
with tab2:
    fig = make_level_chart_fig(df, int(alapszint), selected_class, terulet, angol)
    if fig is None:
        st.info("Nincs megjeleníthető képességszint-adat.")
    else:
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)
with tab3:
    if include_summary:
        fig = make_change_summary_fig(df, selected_class, terulet)
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)
    else:
        st.info("A 3. ábra jelenleg ki van kapcsolva az oldalsávon.")
with tab4:
    if include_distribution:
        fig = make_level_distribution_fig(df, selected_class, terulet, angol)
        if fig is None:
            st.info("Nincs megjeleníthető szintmegoszlás.")
        else:
            st.pyplot(fig, use_container_width=True)
            plt.close(fig)
    else:
        st.info("A 4. ábra jelenleg ki van kapcsolva az oldalsávon.")
with tab5:
    if not include_thresholds:
        st.info("Az 5. ábra jelenleg ki van kapcsolva az oldalsávon.")
    elif thresholds_table.empty:
        st.info("Nincs megjeleníthető ponthatár-tábla ehhez a mérési területhez.")
    else:
        st.dataframe(thresholds_table, use_container_width=True, hide_index=True)

st.subheader("Jelentés letöltése")
docx_bytes = generate_word_report(df, int(alapszint), selected_class, terulet, angol, report_year, previous_year, current_year, institution_name, include_summary, include_distribution, thresholds_table)
st.download_button(
    "📝 Word jelentés letöltése",
    data=docx_bytes,
    file_name=docx_nev if docx_nev.lower().endswith(".docx") else f"{docx_nev}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)

pdf_bytes = generate_pdf(df, int(alapszint), selected_class, terulet, angol, include_summary, include_distribution, thresholds_table)
st.download_button(
    "📄 PDF diagramcsomag letöltése",
    data=pdf_bytes,
    file_name=pdf_nev if pdf_nev.lower().endswith(".pdf") else f"{pdf_nev}.pdf",
    mime="application/pdf",
)
